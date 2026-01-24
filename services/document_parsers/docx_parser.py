# docx_parser.py
# Microsoft Word document (.docx) parser for AURA-NOTES-MANAGER

# Extracts text content, structure (headings, paragraphs), tables, and metadata
# from DOCX files. Provides consistent interface with ParsedDocument output
# compatible with the KG processing pipeline.
#
# Key components:
# - DocxParser: Main parser class with parse() and parse_bytes() methods
# - ParsedDocument: Pydantic model for parsed output
# - DocumentSection: Model for heading-based document sections
# - Table extraction converts to markdown-like format

# @see: api/kg_processor.py - Uses DocxParser for document processing
# @see: AURA-CHAT/backend/document_processor.py - Reference implementation
# @note: Requires python-docx>=0.8.11 library

from __future__ import annotations

import io
import logging
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from pydantic import BaseModel, Field

# Configure module logger
logger = logging.getLogger(__name__)


# =============================================================================
# EXCEPTION CLASSES
# =============================================================================


class DocxParseError(Exception):
    """Base exception for DOCX parsing errors."""

    pass


class CorruptedDocxError(DocxParseError):
    """Raised when DOCX file is corrupted or malformed."""

    pass


class PasswordProtectedDocxError(DocxParseError):
    """Raised when DOCX file is password-protected and cannot be opened."""

    pass


class EmptyDocxError(DocxParseError):
    """Raised when DOCX file contains no text content."""

    pass


# =============================================================================
# PYDANTIC MODELS
# =============================================================================


class DocumentSection(BaseModel):
    """
    Represents a section in a document defined by a heading.

    Attributes:
        level: Heading level (1, 2, 3 for Heading 1, Heading 2, Heading 3)
        title: The heading text
        content: Text content under this heading (until next heading)
        start_position: Character position where section starts in full text
        end_position: Character position where section ends in full text
    """

    level: int = Field(ge=1, le=9, description="Heading level 1-9")
    title: str = Field(description="Heading text")
    content: str = Field(default="", description="Section content text")
    start_position: int = Field(ge=0, description="Start position in document")
    end_position: int = Field(ge=0, description="End position in document")


class ParsedDocument(BaseModel):
    """
    Represents a fully parsed document with text, metadata, and structure.

    Attributes:
        text: Full extracted text content
        metadata: Document metadata (title, author, dates, etc.)
        sections: List of document sections based on headings
        tables: List of tables in markdown-like format
        word_count: Total word count of document
        page_count: Estimated page count (None if unavailable)
    """

    text: str = Field(description="Full extracted text content")
    metadata: Dict[str, Any] = Field(
        default_factory=dict, description="Document metadata"
    )
    sections: List[DocumentSection] = Field(
        default_factory=list, description="Document sections by heading"
    )
    tables: List[str] = Field(
        default_factory=list, description="Tables in markdown format"
    )
    word_count: int = Field(ge=0, description="Total word count")
    page_count: Optional[int] = Field(default=None, description="Estimated page count")


# =============================================================================
# DOCX PARSER CLASS
# =============================================================================


class DocxParser:
    """
    Parser for Microsoft Word .docx files.

    Extracts text content, document structure (headings/sections), tables,
    and metadata from DOCX files. Provides two entry points:
    - parse(file_path): Parse from file path
    - parse_bytes(content): Parse from bytes (for in-memory processing)

    Example:
        parser = DocxParser()
        result = parser.parse("document.docx")
        print(result.text)
        print(result.metadata)
        for section in result.sections:
            print(f"{section.level}: {section.title}")
    """

    # Mapping of python-docx style names to heading levels
    HEADING_STYLES = {
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
        "Heading 5": 5,
        "Heading 6": 6,
        "Heading 7": 7,
        "Heading 8": 8,
        "Heading 9": 9,
        "Title": 0,  # Title is treated as level 0
    }

    def __init__(self):
        """Initialize the DocxParser."""
        self._validate_dependencies()

    def _validate_dependencies(self) -> None:
        """Validate that required dependencies are available."""
        try:
            from docx import Document  # noqa: F401
        except ImportError as e:
            raise ImportError(
                "python-docx library is required. Install with: pip install python-docx"
            ) from e

    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """
        Parse a DOCX file from disk.

        Args:
            file_path: Path to the .docx file

        Returns:
            ParsedDocument with extracted text, metadata, sections, and tables

        Raises:
            FileNotFoundError: If file does not exist
            CorruptedDocxError: If file is corrupted or invalid
            PasswordProtectedDocxError: If file is password-protected
            EmptyDocxError: If file contains no text content
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not file_path.suffix.lower() == ".docx":
            logger.warning(
                f"File {file_path} does not have .docx extension, attempting parse anyway"
            )

        # Check if file is a valid ZIP archive (DOCX is ZIP-based)
        if not zipfile.is_zipfile(file_path):
            raise CorruptedDocxError(
                f"File is not a valid DOCX (ZIP) archive: {file_path}"
            )

        # Check for password protection by examining ZIP contents
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                # Password-protected DOCX files have encrypted content
                # Check for EncryptedPackage which indicates encryption
                if "EncryptedPackage" in zf.namelist():
                    raise PasswordProtectedDocxError(
                        f"File is password-protected: {file_path}"
                    )
        except zipfile.BadZipFile:
            raise CorruptedDocxError(f"File is corrupted (bad ZIP): {file_path}")

        try:
            from docx import Document

            doc = Document(str(file_path))
            return self._parse_document(doc)
        except Exception as e:
            if "password" in str(e).lower() or "encrypted" in str(e).lower():
                raise PasswordProtectedDocxError(
                    f"File is password-protected: {file_path}"
                ) from e
            if "corrupt" in str(e).lower() or "invalid" in str(e).lower():
                raise CorruptedDocxError(f"File is corrupted: {file_path}") from e
            raise DocxParseError(f"Failed to parse DOCX: {e}") from e

    def parse_bytes(self, content: bytes) -> ParsedDocument:
        """
        Parse a DOCX file from bytes.

        Args:
            content: Raw bytes of the .docx file

        Returns:
            ParsedDocument with extracted text, metadata, sections, and tables

        Raises:
            CorruptedDocxError: If content is corrupted or invalid
            PasswordProtectedDocxError: If content is password-protected
            EmptyDocxError: If content contains no text
        """
        if not content:
            raise CorruptedDocxError("Empty content provided")

        # Create BytesIO stream for in-memory processing
        stream = io.BytesIO(content)

        # Check if valid ZIP archive
        if not zipfile.is_zipfile(stream):
            raise CorruptedDocxError("Content is not a valid DOCX (ZIP) archive")

        # Reset stream position after is_zipfile check
        stream.seek(0)

        # Check for password protection
        try:
            with zipfile.ZipFile(stream, "r") as zf:
                if "EncryptedPackage" in zf.namelist():
                    raise PasswordProtectedDocxError("Content is password-protected")
        except zipfile.BadZipFile:
            raise CorruptedDocxError("Content is corrupted (bad ZIP)")

        # Reset stream position for Document parsing
        stream.seek(0)

        try:
            from docx import Document

            doc = Document(stream)
            return self._parse_document(doc)
        except Exception as e:
            if "password" in str(e).lower() or "encrypted" in str(e).lower():
                raise PasswordProtectedDocxError("Content is password-protected") from e
            if "corrupt" in str(e).lower() or "invalid" in str(e).lower():
                raise CorruptedDocxError("Content is corrupted") from e
            raise DocxParseError(f"Failed to parse DOCX bytes: {e}") from e

    def _parse_document(self, doc) -> ParsedDocument:
        """
        Internal method to parse a python-docx Document object.

        Args:
            doc: python-docx Document object

        Returns:
            ParsedDocument with all extracted data
        """
        # Extract components
        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)
        metadata = self._extract_metadata(doc)
        headings = self._extract_headings(doc)

        # Build full text
        full_text = "\n".join(paragraphs)

        # Add table text to full content
        if tables:
            full_text += "\n\n" + "\n\n".join(tables)

        # Check for empty document
        if not full_text.strip():
            raise EmptyDocxError("Document contains no text content")

        # Build sections from headings
        sections = self._build_sections(doc, headings)

        # Calculate word count
        word_count = len(full_text.split())

        # Estimate page count (approximately 500 words per page)
        page_count = max(1, word_count // 500) if word_count > 0 else None

        return ParsedDocument(
            text=full_text,
            metadata=metadata,
            sections=sections,
            tables=tables,
            word_count=word_count,
            page_count=page_count,
        )

    def _extract_paragraphs(self, doc) -> List[str]:
        """
        Extract all paragraph text from document.

        Args:
            doc: python-docx Document object

        Returns:
            List of paragraph text strings
        """
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return paragraphs

    def _extract_tables(self, doc) -> List[str]:
        """
        Extract all tables and convert to markdown-like format.

        Args:
            doc: python-docx Document object

        Returns:
            List of tables as markdown-formatted strings
        """
        tables = []
        for table in doc.tables:
            table_text = self._table_to_markdown(table)
            if table_text:
                tables.append(table_text)
        return tables

    def _table_to_markdown(self, table) -> str:
        """
        Convert a table to markdown-like format.

        Handles merged cells gracefully by using the visible cell text.

        Args:
            table: python-docx Table object

        Returns:
            Markdown-formatted table string
        """
        rows = []
        for row_idx, row in enumerate(table.rows):
            # Extract cell text, handling merged cells
            cells = []
            for cell in row.cells:
                # Get text and clean it
                cell_text = cell.text.strip().replace("\n", " ")
                cells.append(cell_text)

            # Join cells with pipe separator
            row_text = " | ".join(cells)
            rows.append(row_text)

            # Add separator after header row (first row)
            if row_idx == 0:
                separator = " | ".join(["---"] * len(cells))
                rows.append(separator)

        return "\n".join(rows) if rows else ""

    def _extract_metadata(self, doc) -> Dict[str, Any]:
        """
        Extract document metadata (core properties).

        Args:
            doc: python-docx Document object

        Returns:
            Dictionary with metadata fields
        """
        metadata = {}

        try:
            props = doc.core_properties

            # Basic metadata fields
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.subject:
                metadata["subject"] = props.subject
            if props.keywords:
                metadata["keywords"] = props.keywords
            if props.category:
                metadata["category"] = props.category
            if props.comments:
                metadata["comments"] = props.comments

            # Date fields
            if props.created:
                metadata["created"] = self._format_datetime(props.created)
            if props.modified:
                metadata["modified"] = self._format_datetime(props.modified)
            if props.last_printed:
                metadata["last_printed"] = self._format_datetime(props.last_printed)

            # Other fields
            if props.last_modified_by:
                metadata["last_modified_by"] = props.last_modified_by
            if props.revision:
                metadata["revision"] = props.revision
            if props.version:
                metadata["version"] = props.version

        except Exception as e:
            logger.warning(f"Failed to extract some metadata: {e}")

        return metadata

    def _format_datetime(self, dt: Optional[datetime]) -> str:
        """Format datetime to ISO string, handling None values."""
        if dt is None:
            return ""
        try:
            return dt.isoformat()
        except Exception:
            return str(dt)

    def _extract_headings(self, doc) -> List[Tuple[int, str, int]]:
        """
        Extract all headings with their levels and positions.

        Args:
            doc: python-docx Document object

        Returns:
            List of tuples: (level, title, paragraph_index)
        """
        headings = []
        for idx, para in enumerate(doc.paragraphs):
            if para.style and para.style.name in self.HEADING_STYLES:
                level = self.HEADING_STYLES[para.style.name]
                title = para.text.strip()
                if title:  # Only add non-empty headings
                    headings.append((level, title, idx))
        return headings

    def _build_sections(
        self, doc, headings: List[Tuple[int, str, int]]
    ) -> List[DocumentSection]:
        """
        Build document sections from headings.

        Each section includes the heading and all content until the next heading.

        Args:
            doc: python-docx Document object
            headings: List of (level, title, paragraph_index) tuples

        Returns:
            List of DocumentSection objects
        """
        sections = []
        paragraphs = doc.paragraphs

        if not headings:
            return sections

        # Build full text for position tracking
        full_text_parts = []
        para_positions = []  # (start, end) for each paragraph
        current_pos = 0

        for para in paragraphs:
            text = para.text
            start = current_pos
            end = current_pos + len(text)
            para_positions.append((start, end))
            full_text_parts.append(text)
            current_pos = end + 1  # +1 for newline

        # Build sections
        for i, (level, title, para_idx) in enumerate(headings):
            # Determine content range
            content_start_idx = para_idx + 1  # Start after heading

            # End at next heading or document end
            if i + 1 < len(headings):
                content_end_idx = headings[i + 1][2]
            else:
                content_end_idx = len(paragraphs)

            # Collect content paragraphs
            content_paragraphs = []
            for j in range(content_start_idx, content_end_idx):
                if j < len(paragraphs):
                    text = paragraphs[j].text.strip()
                    if text:
                        content_paragraphs.append(text)

            content = "\n".join(content_paragraphs)

            # Calculate positions
            start_pos = (
                para_positions[para_idx][0] if para_idx < len(para_positions) else 0
            )
            if content_end_idx > 0 and content_end_idx <= len(para_positions):
                end_pos = para_positions[content_end_idx - 1][1]
            else:
                end_pos = start_pos + len(title) + len(content)

            sections.append(
                DocumentSection(
                    level=max(1, level),  # Ensure minimum level of 1
                    title=title,
                    content=content,
                    start_position=start_pos,
                    end_position=end_pos,
                )
            )

        return sections
